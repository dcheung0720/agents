import streamlit as st
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

st.title("DOCX Viewer with Preserved Order")

uploaded_file = st.file_uploader("Upload a .docx file", type="docx")

if uploaded_file is not None:
    doc = Document(uploaded_file)
    html_content = ""

    # Helper: iterate through elements in order
    def iter_block_items(parent):
        for child in parent.element.body:
            if isinstance(child, CT_P):
                yield doc.paragraphs[len(list(child.xpath("preceding-sibling::w:p")))]
            elif isinstance(child, CT_Tbl):
                yield doc.tables[len(list(child.xpath("preceding-sibling::w:tbl")))]

    # Build HTML
    for block in iter_block_items(doc):
        if isinstance(block, type(doc.paragraphs[0])):  # Paragraph
            if block.text.strip():
                html_content += f"<p>{block.text}</p>"
        else:  # Table
            html_content += "<table border='1' style='border-collapse:collapse; width:100%; margin-top:10px;'>"
            for row in block.rows:
                html_content += "<tr>"
                for cell in row.cells:
                    html_content += f"<td style='padding:5px;'>{cell.text}</td>"
                html_content += "</tr>"
            html_content += "</table>"

    st.markdown(html_content, unsafe_allow_html=True)
